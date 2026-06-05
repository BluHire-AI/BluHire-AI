import pdfplumber
import fitz  # PyMuPDF
from docx import Document
import io
import os

class ResumeParser:
    @staticmethod
    def extract_text_from_pdf(file_bytes: bytes) -> str:
        text = ""
        # Try pdfplumber first
        try:
            with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
        except Exception as e:
            print(f"pdfplumber extraction failed: {e}")

        # Fallback to PyMuPDF (fitz) if empty
        if not text.strip():
            try:
                doc = fitz.open(stream=file_bytes, filetype="pdf")
                for page in doc:
                    page_text = page.get_text()
                    if page_text:
                        text += page_text + "\n"
            except Exception as e:
                print(f"PyMuPDF extraction failed: {e}")

        return text.strip()

    @staticmethod
    def extract_text_from_docx(file_bytes: bytes) -> str:
        text = ""
        try:
            doc = Document(io.BytesIO(file_bytes))
            for para in doc.paragraphs:
                if para.text:
                    text += para.text + "\n"
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
        except Exception as e:
            print(f"python-docx extraction failed: {e}")
        return text.strip()

    @classmethod
    def parse(cls, filename: str, file_bytes: bytes) -> str:
        _, ext = os.path.splitext(filename.lower())
        if ext == '.pdf':
            return cls.extract_text_from_pdf(file_bytes)
        elif ext in ['.docx', '.doc']:
            return cls.extract_text_from_docx(file_bytes)
        else:
            try:
                return file_bytes.decode('utf-8', errors='ignore')
            except Exception:
                return ""
